# fix_captions.py
import sys

from docx import Document
from lxml import etree

NS  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WNS = f'{{{NS}}}'
CAPTION_STYLES = {'ImageCaption', 'TableCaption'}

def fix_caption_in_cell(cell):
    # 单元格内若含嵌套 <w:tbl> → 表格语境，否则 → 图片语境
    has_nested_tbl = bool(cell._tc.findall(f'{WNS}tbl'))
    target_style   = 'TableCaption' if has_nested_tbl else 'ImageCaption'
    count = 0

    for para in cell.paragraphs:
        p       = para._p
        pPr_all = p.findall(f'{WNS}pPr')

        # 找到含题注样式的 pPr
        main_pPr = next(
            (pPr for pPr in pPr_all
             if (s := pPr.find(f'{WNS}pStyle')) is not None
             and s.get(f'{WNS}val') in CAPTION_STYLES),
            None
        )
        if main_pPr is None:
            continue

        # 删除多余的重复 pPr（兼容图片题注残留的双 pPr 问题）
        for pPr in pPr_all:
            if pPr is not main_pPr:
                p.remove(pPr)

        # 按语境写入正确样式
        main_pPr.find(f'{WNS}pStyle').set(f'{WNS}val', target_style)

        # 删除旧 jc，写入居中
        for jc in main_pPr.findall(f'{WNS}jc'):
            main_pPr.remove(jc)
        jc_elem = etree.SubElement(main_pPr, f'{WNS}jc')
        jc_elem.set(f'{WNS}val', 'center')

        count += 1

    return count

def fix_all(path):
    doc   = Document(path)
    count = 0
    for table in doc.tables:          # 只遍历顶层表格，嵌套数据表不会被访问
        for row in table.rows:
            for cell in row.cells:
                count += fix_caption_in_cell(cell)
    doc.save(path)
    print(f"修正了 {count} 处题注样式和对齐")

fix_all(sys.argv[1])